from docx2python import docx2python
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import re
from .footnote import Footnote

class FootnoteParser:

	def __init__(self):
		self.jurisprudences = []
		self.monographs = []
		self.legislations = []
		self.journals = []
		self.online = []
		self.others = []

		self._parsers = {"juris": self._parse_jurisprudences,
		                "legis": self._parse_legislations,
						"mono": self._parse_monographs,
						"journo": self._parse_journals,
						"online": self._parse_onlines}
		
		self._patterns = {
			"juris": r'^(.*) ?((?:\[|\().*)', # (R v Nanabush ) ([1999] rest of citation)
			"legis": r'^(.+?),(.*)$', # (Optional bill number, title), (rest of citation)
			"author_1": r'^(.+?),(.*)$',
			"author_2": r'(.+?) (?:\&|and) (.+?), (.*)',
			"author_3": r'(.+?),(.+?),(.*)',
			"authors_4": r'(.+?, et.? al.?), (.*)'
		}

	def _get_footnotes_from_doc(self, src):
		doc = docx2python(src)

		for note in doc.footnotes[0][0][0]:
			note = note.split(";")
			for n in note:
				match = re.search(r'{f (juris|legis|mono|journo|online)(?: (\d+))?}(.*)(?:{\/f})', n.strip())
				if match:
					self._parse_types(match.groups())

	def _parse_types(self, groups):
		self._parsers[groups[0]](groups[1], groups[2])

	def _parse_jurisprudences(self, _, citation):
		# first arg is ignored because uniform legal citation is pretty standard
		match = re.match(self._patterns["juris"], citation)
		if match:
			groups = match.groups()
			self.jurisprudences.append(Footnote(groups[0], None, groups[1]))
		else:
			print("No match found while parsing jurisprudence: %s..." % citation[0:20])


	def _parse_legislations(self, _, citation):
		# same as above, first arg already ignored
		match = re.match(self._patterns["legis"], citation)
		if match:
			groups = match.groups()
			self.legislations.append(Footnote(groups[0], None, groups[1]))
		else:
			print("No match found while parsing legislation: %s..." % citation[0:20])

	def _parse_monographs(self, authors, citation):
		match = re.match(self._patterns["author_%s" % authors], citation)
		if match:
			groups = match.groups()
			switched = self._switch_author_name(groups[0])

			if (len(groups) == 3):
				self.monographs.append(Footnote(switched, groups[1], groups[2]))
			else:
				self.monographs.append(Footnote(switched, None, groups[1]))
		else:
			print("No match was found while parsing monograph: %s..." % citation[0:20])

	def _parse_journals(self, authors, citation):
		match = re.match(self._patterns["author_%s" % authors], citation)
		if match:
			groups = match.groups()
			switched = self._switch_author_name(groups[0])

			if (len(groups) == 3):
				self.journals.append(Footnote(switched, groups[1], groups[2]))
			else:
				self.journals.append(Footnote(switched, None, groups[1]))
		else:
			print("No match was found while parsing monograph: %s..." % citation[0:20])

	def _parse_onlines(self, authors, citation):
		match = re.match(self._patterns["author_%s" % authors], citation)
		if match:
			groups = match.groups()
			switched = self._switch_author_name(groups[0])

			if (len(groups) == 3):
				self.online.append(Footnote(switched, groups[1], groups[2]))
			else:
				self.online.append(Footnote(switched, None, groups[1]))
		else:
			print("No match was found while parsing monograph: %s..." % citation[0:20])

	def _switch_author_name(self, string):
		match = re.match(r'(.*) (.*)', string)
		return match.groups()[1] + ", " + match.groups()[0]

	def sortall(self):
		self.legislations.sort(key=lambda note: note.lead)
		self.jurisprudences.sort(key=lambda note: note.lead)
		self.journals.sort(key=lambda note: note.lead)
		self.monographs.sort(key=lambda note: note.lead)
		self.online.sort(key=lambda note: note.lead)

	def _write_to_doc(self, dst):
		doc = Document()
		doc.add_heading("Bibliography", 0)

		self._create_source_header(doc, "Legislation")
		for l in self.legislations:
			self._create_source_paragraph(doc, l.__str__())
		
		self._create_source_header(doc, "Jurisprudence")
		for j in self.jurisprudences:
			self._create_source_paragraph(doc, j.__str__())
		
		self._create_source_header(doc, "Secondary Sources: Monographs")
		for m in self.monographs:
			self._create_source_paragraph(doc, m.__str__())
		
		self._create_source_header(doc, "Secondary Sources: Journals/Articles")
		for jo in self.journals:
			self._create_source_paragraph(doc, jo.__str__())
		
		self._create_source_header(doc, "Secondary Sources: Online")
		for o in self.online:
			self._create_source_paragraph(doc, o.__str__())
		
		doc.save(dst)

	def _create_source_header(self, doc, text):
		p = doc.add_heading("", level=1)
		p.paragraph_format.space_before = Pt(0)
		p.paragraph_format.space_before = Pt(0)
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
		r = p.add_run(text)
		r.font.size = Pt(12)
		r.font.all_caps = True
		r.font.bold = True
	
	def _create_source_paragraph(self, doc, text):
		p = doc.add_paragraph(text)
		p.paragraph_format.space_before = Pt(0)
		p.paragraph_format.space_after = Pt(0)
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

	def make_biblio(self, src_path, dst_path):
		self._get_footnotes_from_doc(src_path)
		self.sortall()
		self._write_to_doc(dst_path)
